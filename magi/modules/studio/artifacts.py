"""
Fábrica de artefactos con bucle de observación (Plan MAGI 9.0 §5).

EL PATRÓN
=========
    ESPECIFICAR -> GENERAR -> EJECUTAR/RENDERIZAR -> OBSERVAR -> CRITICAR -> ITERAR

La clave es OBSERVAR. Un sistema que genera un juego y te lo entrega sin
haberlo arrancado ha *generado código de juego*; uno que lo arranca, captura un
fotograma y lo mira, ha *hecho un juego*. La diferencia no está en el modelo:
está en si hay un bucle que cierra sobre el resultado.

Es la misma idea que la verificación ejecutable del §2.5, aplicada a artefactos
que no son código: una imagen se mira, un documento se renderiza y se cuenta
las páginas, un juego se arranca en modo headless y se captura.
"""
from __future__ import annotations

import asyncio
import logging
import shutil
import subprocess
import sys
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


class ArtifactKind(str, Enum):
    PROGRAM = "programa"
    GAME = "juego"
    IMAGE = "imagen"
    DOCUMENT = "documento"
    VIDEO = "video"
    DATA = "datos"


@dataclass
class Observation:
    """Lo que el sistema VE de su propio resultado."""
    ok: bool
    kind: ArtifactKind
    summary: str
    evidence: list[str] = field(default_factory=list)
    artifact_path: str | None = None
    screenshot: str | None = None
    problems: list[str] = field(default_factory=list)

    def render(self) -> str:
        head = f"[{'OK' if self.ok else 'FALLA'}] {self.kind.value}: {self.summary}"
        parts = [head]
        if self.evidence:
            parts += [f"  · {e}" for e in self.evidence]
        if self.problems:
            parts.append("  problemas observados:")
            parts += [f"  · {p}" for p in self.problems]
        if self.screenshot:
            parts.append(f"  captura: {self.screenshot}")
        return "\n".join(parts)

    def feedback(self) -> str:
        """Lo que vuelve al generador cuando algo no cuadra."""
        if self.ok and not self.problems:
            return ""
        return ("El artefacto se ha inspeccionado y presenta estos problemas. "
                "Corrígelos:\n" + "\n".join(f"- {p}" for p in self.problems))


# --------------------------------------------------------------- programas

async def observe_program(path: str | Path, *, entry: str = "",
                          timeout: int = 60) -> Observation:
    """Arranca el programa y mira si sobrevive."""
    p = Path(path)
    if not p.exists():
        return Observation(False, ArtifactKind.PROGRAM, "no existe",
                           problems=[f"{p} no existe"])

    cmd = entry or f'"{sys.executable}" "{p.name}"'
    cwd = p if p.is_dir() else p.parent
    rc, out = await _run(cmd, cwd, timeout)

    problems = []
    if rc != 0:
        problems.append(f"termina con código {rc}")
        tail = "\n".join(out.strip().splitlines()[-6:])
        if tail:
            problems.append(f"últimas líneas:\n{tail}")
    return Observation(rc == 0, ArtifactKind.PROGRAM,
                       f"ejecutado (rc={rc})",
                       evidence=[out[-800:]] if out else [],
                       artifact_path=str(p), problems=problems)


# ------------------------------------------------------------------ juegos

PYGAME_HARNESS = '''"""
Arnés de observación generado por MAGI (§5.2).

Arranca el juego en modo headless, avanza N fotogramas y guarda una captura.
Sin esto el sistema entrega "código de juego" sin saber si el jugador se
distingue del fondo, o si la pantalla sale en negro.
"""
import os, sys, importlib.util
os.environ.setdefault("SDL_VIDEODRIVER", "dummy")
os.environ.setdefault("SDL_AUDIODRIVER", "dummy")

import pygame

FRAMES = int(os.environ.get("MAGI_FRAMES", "120"))
SHOT = os.environ.get("MAGI_SHOT", "frame.png")
TARGET = os.environ.get("MAGI_TARGET", "main.py")

pygame.init()
_orig_flip, _orig_update = pygame.display.flip, pygame.display.update
state = {"frames": 0}


def _capture():
    surf = pygame.display.get_surface()
    if surf is not None:
        pygame.image.save(surf, SHOT)


def _flip(*a, **kw):
    state["frames"] += 1
    if state["frames"] >= FRAMES:
        _capture()
        pygame.quit()
        raise SystemExit(0)
    return _orig_flip(*a, **kw)


pygame.display.flip = _flip
pygame.display.update = _flip

spec = importlib.util.spec_from_file_location("__main__", TARGET)
mod = importlib.util.module_from_spec(spec)
try:
    spec.loader.exec_module(mod)
except SystemExit:
    pass
finally:
    if state["frames"] > 0:
        _capture()
    print(f"MAGI_FRAMES_RENDERED={state['frames']}")
'''


async def observe_game(project_dir: str | Path, *, entry: str = "main.py",
                       frames: int = 120, timeout: int = 90) -> Observation:
    """
    Arranca un juego Pygame en headless, avanza fotogramas y captura uno.

    Lo que esto permite: que Balthasar MIRE la captura con visión y diga "el
    jugador no se distingue del fondo" — una crítica imposible de hacer
    leyendo el código.
    """
    d = Path(project_dir)
    target = d / entry
    if not target.exists():
        return Observation(False, ArtifactKind.GAME, "sin punto de entrada",
                           problems=[f"no existe {target}"])

    try:
        import pygame  # noqa: F401
    except ImportError:
        return Observation(
            False, ArtifactKind.GAME, "pygame no instalado",
            problems=["pip install pygame para poder observar el juego. "
                      "Sin esto solo se puede revisar el código, no verlo."])

    harness = d / "_magi_harness.py"
    harness.write_text(PYGAME_HARNESS, encoding="utf-8")
    shot = d / "_magi_frame.png"
    # Borrar la captura anterior: si el juego revienta antes de dibujar, el
    # fichero viejo sigue ahí y el informe describe la ejecución EQUIVOCADA.
    shot.unlink(missing_ok=True)
    env = {"MAGI_FRAMES": str(frames), "MAGI_SHOT": str(shot),
           "MAGI_TARGET": str(target), "SDL_VIDEODRIVER": "dummy",
           "SDL_AUDIODRIVER": "dummy"}

    try:
        rc, out = await _run(f'"{sys.executable}" "{harness.name}"', d,
                             timeout, env)
    finally:
        harness.unlink(missing_ok=True)

    rendered = 0
    for line in out.splitlines():
        if line.startswith("MAGI_FRAMES_RENDERED="):
            rendered = int(line.split("=")[1] or 0)

    problems = []
    if rendered == 0:
        problems.append("no se dibujó ni un fotograma: el juego no llega a "
                        "arrancar o nunca llama a display.flip()")
    if not shot.exists():
        problems.append("no se pudo capturar la pantalla")
    if rc != 0 and rendered == 0:
        problems.append(f"salida con código {rc}:\n" +
                        "\n".join(out.strip().splitlines()[-6:]))

    evidence = [f"código de salida {rc}"]
    if shot.exists():
        desc = _describe_image(shot)
        evidence.append(desc)
        # El fallo que este bucle existe para cazar: el juego corre, dibuja
        # fotogramas y en pantalla no se ve nada porque todo es del mismo
        # color. Sin esto el informe decía OK y enterraba la pista en la
        # evidencia.
        if "VACÍA" in desc:
            problems.append(
                "la pantalla es de un solo color: el juego dibuja pero no se "
                "ve nada. Revisa que los elementos no sean del color del fondo "
                "y que se dibujen dentro de los límites de la ventana.")

    return Observation(
        bool(rendered) and shot.exists() and not problems, ArtifactKind.GAME,
        f"{rendered} fotogramas dibujados",
        evidence=evidence, artifact_path=str(d),
        screenshot=str(shot) if shot.exists() else None,
        problems=problems)


# ------------------------------------------------------------------ imagen

def _describe_image(path: str | Path) -> str:
    """
    Descripción objetiva de una imagen SIN modelo de visión.

    Detecta el fallo más común de una captura de juego: la pantalla en negro,
    o una imagen de un solo color. Es barato y no gasta cuota.
    """
    try:
        from PIL import Image
    except ImportError:
        return "Pillow no instalado: no se puede inspeccionar la imagen"
    try:
        with Image.open(path) as im:
            im = im.convert("RGB")
            w, h = im.size
            colors = im.getcolors(maxcolors=1_000_000)
            if not colors:
                return f"{w}x{h}, muchos colores distintos"
            colors.sort(reverse=True)
            top_count, top_rgb = colors[0]
            share = top_count / (w * h)
            desc = (f"{w}x{h}, {len(colors)} colores; el dominante "
                    f"{top_rgb} ocupa el {share:.0%}")
            if share > 0.98:
                desc += "  <-- PANTALLA PRÁCTICAMENTE VACÍA"
            return desc
    except Exception as e:
        return f"imagen ilegible: {e}"


async def observe_image(path: str | Path) -> Observation:
    p = Path(path)
    if not p.exists():
        return Observation(False, ArtifactKind.IMAGE, "no existe",
                           problems=[f"{p} no existe"])
    desc = _describe_image(p)
    problems = []
    if "VACÍA" in desc:
        problems.append("la imagen es casi de un solo color: probablemente no "
                        "se dibujó nada")
    if "ilegible" in desc:
        problems.append(desc)
    return Observation(not problems, ArtifactKind.IMAGE, desc,
                       artifact_path=str(p), screenshot=str(p),
                       problems=problems)


# -------------------------------------------------------------- documentos

async def observe_document(path: str | Path) -> Observation:
    """
    Cuenta páginas, palabras y detecta documentos vacíos.

    El fallo típico de generación de documentos es entregar un .docx con la
    plantilla y sin contenido, o un PDF de una sola página cuando se pidieron
    veinte.
    """
    p = Path(path)
    if not p.exists():
        return Observation(False, ArtifactKind.DOCUMENT, "no existe",
                           problems=[f"{p} no existe"])

    ext = p.suffix.lower()
    evidence, problems = [f"{p.stat().st_size:,} bytes"], []

    if ext == ".pdf":
        try:
            import pypdf
            r = pypdf.PdfReader(str(p))
            n = len(r.pages)
            text = "".join((pg.extract_text() or "") for pg in r.pages[:20])
            evidence.append(f"{n} páginas, {len(text.split())} palabras "
                            f"en las primeras {min(n, 20)}")
            if n == 0:
                problems.append("PDF sin páginas")
            elif len(text.strip()) < 40:
                problems.append("el PDF no tiene texto extraíble: ¿páginas en "
                                "blanco o solo imágenes?")
        except ImportError:
            evidence.append("pypdf no instalado: no se puede inspeccionar")
        except Exception as e:
            problems.append(f"PDF ilegible: {e}")

    elif ext in (".docx", ".dotx"):
        try:
            import docx
            d = docx.Document(str(p))
            words = sum(len(par.text.split()) for par in d.paragraphs)
            evidence.append(f"{len(d.paragraphs)} párrafos, {words} palabras, "
                            f"{len(d.tables)} tablas")
            if words < 20:
                problems.append("el documento está prácticamente vacío")
        except ImportError:
            evidence.append("python-docx no instalado")
        except Exception as e:
            problems.append(f"docx ilegible: {e}")

    elif ext in (".md", ".txt", ".html"):
        text = p.read_text(encoding="utf-8", errors="replace")
        words = len(text.split())
        evidence.append(f"{words} palabras, {len(text.splitlines())} líneas")
        if words < 20:
            problems.append("prácticamente vacío")

    else:
        evidence.append(f"tipo {ext or 'sin extensión'} no inspeccionable")

    if p.stat().st_size == 0:
        problems.append("fichero de 0 bytes")

    return Observation(not problems, ArtifactKind.DOCUMENT,
                       f"{p.name} ({ext or 'sin extensión'})",
                       evidence=evidence, artifact_path=str(p),
                       problems=problems)


# ----------------------------------------------------------------- común

async def _run(cmd: str, cwd: Path, timeout: int,
               extra_env: dict | None = None) -> tuple[int, str]:
    import os
    env = {**os.environ, **(extra_env or {})}
    cwd.mkdir(parents=True, exist_ok=True)
    try:
        proc = await asyncio.create_subprocess_shell(
            cmd, cwd=str(cwd), env=env,
            stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.STDOUT)
    except Exception as e:
        return 127, str(e)
    try:
        out, _ = await asyncio.wait_for(proc.communicate(), timeout=timeout)
    except asyncio.TimeoutError:
        proc.kill()
        await proc.wait()
        return 124, f"timeout tras {timeout}s"
    return proc.returncode or 0, out.decode("utf-8", errors="replace")


async def observe(path: str | Path, kind: ArtifactKind | str | None = None,
                  **kw) -> Observation:
    """Despacha por tipo, deduciéndolo de la extensión si no se indica."""
    p = Path(path)
    if kind is None:
        ext = p.suffix.lower()
        if p.is_dir():
            kind = ArtifactKind.GAME
        elif ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp"):
            kind = ArtifactKind.IMAGE
        elif ext in (".pdf", ".docx", ".md", ".txt", ".html", ".dotx"):
            kind = ArtifactKind.DOCUMENT
        else:
            kind = ArtifactKind.PROGRAM
    kind = ArtifactKind(kind) if isinstance(kind, str) else kind

    if kind is ArtifactKind.GAME:
        return await observe_game(p, **kw)
    if kind is ArtifactKind.IMAGE:
        return await observe_image(p)
    if kind is ArtifactKind.DOCUMENT:
        return await observe_document(p)
    return await observe_program(p, **kw)


def available_backends() -> dict[str, bool]:
    """Qué se puede observar en esta máquina."""
    def has(mod: str) -> bool:
        import importlib.util
        return importlib.util.find_spec(mod) is not None

    return {
        "pygame": has("pygame"), "pillow": has("PIL"),
        "pypdf": has("pypdf"), "python-docx": has("docx"),
        "ffmpeg": bool(shutil.which("ffmpeg")),
        "comfyui_local": _comfy_reachable(),
    }


def _comfy_reachable(host: str = "http://127.0.0.1:8188") -> bool:
    """ComfyUI para imagen/manga (§5.4). Local, gratis, sin claves."""
    try:
        import urllib.request
        with urllib.request.urlopen(f"{host}/system_stats", timeout=1):
            return True
    except Exception:
        return False
